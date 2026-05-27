"""
Generate comprehensive technical documentation for the Community Agent Town project.
Writes to Generative_Agents_Project_Documentation.docx.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── helpers ──────────────────────────────────────────────────────────────────

def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_para(doc, text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p

def add_code(doc, code_text):
    """Add a monospace code block."""
    style = doc.styles["No Spacing"]
    p = doc.add_paragraph(style=style)
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0x6E)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p

def add_numbered(doc, text):
    doc.add_paragraph(text, style="List Number")

def section_break(doc):
    doc.add_paragraph()

# ── main ─────────────────────────────────────────────────────────────────────

def build_doc():
    doc = Document()

    # Title
    title = doc.add_heading("Community Agent Town — Technical Documentation", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Version 0.2  |  Generated {datetime.date.today()}  |  For junior developers").italic = True

    doc.add_page_break()

    # ── TABLE OF CONTENTS (manual) ────────────────────────────────────────────
    add_heading(doc, "Table of Contents", 1)
    toc_items = [
        "1. Project Overview",
        "2. Repository Structure",
        "3. Architecture Diagram",
        "4. Core Algorithms",
        "   4.1 Agent Tick Lifecycle",
        "   4.2 Memory Retrieval Scoring",
        "   4.3 Action Resolution Pipeline",
        "   4.4 Importance Scoring",
        "5. Line-by-Line Code Analysis",
        "   5.1 src/agents/agent.py — Agent class",
        "   5.2 src/engine/simulation.py — SimulationEngine",
        "   5.3 src/cognition/memory.py — MemoryService",
        "   5.4 src/world/resolver.py — ActionResolver",
        "   5.5 src/prompts/act.py — Action prompt",
        "6. Data Flow Diagrams (ASCII Flowcharts)",
        "   6.1 Full Tick Flow",
        "   6.2 Memory Retrieval Flow",
        "   6.3 Proposal → Resolution Flow",
        "7. Database Schema",
        "8. Configuration Reference",
        "9. Source Code Change History",
        "10. How to Run & Test",
        "11. Guardrails & Constraints",
        "12. Glossary",
    ]
    for item in toc_items:
        add_para(doc, item)

    doc.add_page_break()

    # ── 1. PROJECT OVERVIEW ───────────────────────────────────────────────────
    add_heading(doc, "1. Project Overview", 1)
    add_para(doc, (
        "Community Agent Town is a multi-agent simulation in which three AI-driven characters — "
        "Maria (cafe owner), John (novelist), and Emma (folklore researcher) — live out daily routines "
        "inside a small fictional town. Each agent perceives its environment, reflects on experiences, "
        "creates daily plans, and acts autonomously using a local large-language model (LLM). "
        "All behaviour is grounded: agents may only visit real locations, interact with real objects, "
        "and speak to agents who are physically present."
    ))
    section_break(doc)

    add_para(doc, "Key design goals:", bold=True)
    goals = [
        "Emergent, believable behaviour — agents develop relationships and routines without scripting.",
        "Strict groundedness — the LLM is never allowed to invent people, locations, crimes, or historical facts.",
        "Transparency — every decision is logged to a transcript, event JSONL, and SQLite database.",
        "Testability — a deterministic harness allows regression testing without a live LLM.",
        "Extensibility — new locations, objects, and agents can be added via YAML config files.",
    ]
    for g in goals:
        add_bullet(doc, g)

    doc.add_page_break()

    # ── 2. REPOSITORY STRUCTURE ───────────────────────────────────────────────
    add_heading(doc, "2. Repository Structure", 1)
    add_para(doc, "Every file and folder with its purpose:", bold=True)

    structure = """
comunity Agent town/
├── agents/                     Runtime markdown files (created at startup)
│   ├── emma/
│   │   ├── SOUL.md             Stable personality & guardrails (hand-editable)
│   │   ├── KNOWLEDGE.md        Durable facts appended during reflections
│   │   └── TODAY.md            Daily schedule, reset at 08:00 each sim-day
│   ├── john/  (same structure)
│   └── maria/ (same structure)
│
├── config/
│   ├── agents.yaml             Agent definitions: name, bio, goals, traits
│   ├── simulation.yaml         Tick duration, reflection threshold, etc.
│   └── world.yaml              Locations, objects, affordances, allowed states
│
├── src/
│   ├── main.py                 Entry point (CLI: --ticks, --start-llama-server, …)
│   ├── agents/
│   │   ├── agent.py            Core Agent class (perceive→reflect→plan→act)
│   │   ├── files.py            Reads/writes SOUL/KNOWLEDGE/TODAY.md
│   │   ├── personality.py      Builds character capsule for prompts
│   │   ├── proposals.py        Data classes: MoveProposal, SpeakProposal, …
│   │   └── soul.py             Ensures SOUL.md exists on disk
│   ├── cognition/
│   │   ├── memory.py           MemoryService: add, retrieve, semantic memory
│   │   ├── embeddings.py       Embedding abstraction (real vs fake)
│   │   ├── needs.py            AgentNeeds: hunger, energy, social_satiety
│   │   ├── planner.py          Planning service (placeholder)
│   │   └── reflection.py       Reflection helpers
│   ├── engine/
│   │   ├── simulation.py       SimulationEngine: tick loop orchestration
│   │   ├── clock.py            Sim-time tracking
│   │   └── replay.py           Event replay logic
│   ├── harness/
│   │   └── agent_harness.py    Deterministic smoke-test (fake LLM + embeddings)
│   ├── llm/
│   │   ├── client.py           LlamaCppClient — OpenAI-compatible HTTP client
│   │   ├── gateway.py          LLMGateway — routes ACT/PLAN/REFLECT calls
│   │   ├── benchmark.py        GPU verification & llama.cpp server management
│   │   └── parsing.py          JSON response parsing + fallback
│   ├── observability/
│   │   ├── event_log.py        Writes events.jsonl
│   │   └── transcript.py       Human-readable simulation.log
│   ├── prompts/
│   │   ├── act.py              Action selection prompt (12 000 char budget)
│   │   ├── plan.py             Daily schedule prompt
│   │   ├── reflect.py          Reflection & knowledge-extraction prompt
│   │   ├── dialogue.py         (placeholder)
│   │   ├── perceive.py         (placeholder)
│   │   └── score_importance.py Rule-based importance scoring
│   ├── storage/
│   │   ├── db.py               SQLite / Postgres connection factory
│   │   ├── schema.sql          SQLite table definitions
│   │   ├── postgres_schema.sql Postgres table definitions
│   │   └── repositories.py     Async data-access layer
│   └── world/
│       ├── state.py            WorldState: positions, objects, locations
│       ├── context.py          Builds action menu for agents
│       ├── events.py           Event data classes (Move, Speech, …)
│       └── resolver.py         ActionResolver: validates & applies proposals
│
├── tests/                      pytest unit & integration tests
├── pyproject.toml              Package config & dependencies
└── README.md                   Setup guide, llama.cpp config, run commands
""".strip()
    add_code(doc, structure)

    doc.add_page_break()

    # ── 3. ARCHITECTURE DIAGRAM ───────────────────────────────────────────────
    add_heading(doc, "3. Architecture Diagram", 1)
    add_para(doc, (
        "The diagram below shows how the major components relate. "
        "Arrows represent data or control flow."
    ))

    arch = """
┌──────────────────────────────────────────────────────────────────────┐
│                           src/main.py                                │
│  Parses CLI args → loads config YAML → creates DB / LLM / World     │
│  → creates Agents → runs SimulationEngine.run(num_ticks)             │
└─────────────────────┬────────────────────────────────────────────────┘
                      │ calls tick() N times
                      ▼
┌──────────────────────────────────────────────────────────────────────┐
│                     SimulationEngine  (engine/simulation.py)         │
│                                                                      │
│  ① daily_heartbeat()  — reset plan & files at 08:00                 │
│  ② perceive()         — parallel across all agents                  │
│  ③ maybe_reflect()    — sequential (importance gate)                │
│  ④ maybe_plan()       — sequential (once per day or on demand)      │
│  ⑤ end_of_day_heartbeat() — consolidate memories at 22:00          │
│  ⑥ propose_action()  — parallel across all agents                  │
│  ⑦ resolve()         — sorted by priority (Speak/Use > Move > Wait) │
└──────────┬──────────────────────────────────────────────────────────┘
           │ shared references
     ┌─────┴──────┐
     │            │
     ▼            ▼
┌─────────────────────┐    ┌─────────────────────────────────────────────┐
│   WorldState        │    │         Agent  (agents/agent.py)            │
│ - agent_positions   │    │                                             │
│ - locations (YAML)  │◄───│  perceive()  — deterministic observations   │
│ - objects + states  │    │  reflect()   — LLM: insights + knowledge    │
└─────────────────────┘    │  plan()      — LLM: hour-by-hour schedule   │
                           │  propose()   — LLM: pick action from menu   │
┌─────────────────────┐    │  _proposal_from_dict() — parse LLM output   │
│   EventBus          │◄───│  _avoid_stale_action() — anti-loop guard    │
│ - publish/subscribe │    └──────────────┬──────────────────────────────┘
│ - events per tick   │                   │ uses
└─────────────────────┘             ┌─────┴──────────────────┐
                                    │   LLMGateway           │
┌─────────────────────┐             │  gateway.py            │
│   MemoryService     │             │  Routes ACT/PLAN/REFLECT│
│  cognition/memory.py│             │  → LlamaCppClient      │
│  - add()            │             │  (localhost:8080)      │
│  - retrieve()       │◄────────────└────────────────────────┘
│  - add_semantic()   │
└─────────────────────┘
         │ persists to
         ▼
┌─────────────────────────────────────────────────────────────────────┐
│   SQLite / Postgres  (storage/repositories.py)                      │
│   Tables: agents, memories, plans, relationships,                    │
│           semantic_memory, events, call_logs                         │
└─────────────────────────────────────────────────────────────────────┘
         │ also writes to
         ▼
   events.jsonl   simulation.log   agents/*/KNOWLEDGE.md   agents/*/TODAY.md
""".strip()
    add_code(doc, arch)

    doc.add_page_break()

    # ── 4. CORE ALGORITHMS ────────────────────────────────────────────────────
    add_heading(doc, "4. Core Algorithms", 1)

    # 4.1 Agent Tick Lifecycle
    add_heading(doc, "4.1  Agent Tick Lifecycle", 2)
    add_para(doc, (
        "Every simulated tick (= 30 sim-minutes) each agent executes five steps in order. "
        "Steps 1-4 may involve LLM calls; step 5 is purely rule-based."
    ))
    steps = [
        ("perceive()", (
            "Build up to 3 textual observations from: (a) a fallback 'I am at X with Y' sentence, "
            "(b) events that occurred in this location last tick (speech, moves, object changes, rejections), "
            "(c) a summary of object states nearby. Each observation is importance-scored (rule-based, 1–10) "
            "and stored in the memory database. The importance score is added to importance_accumulator."
        )),
        ("maybe_reflect()", (
            "Fires only when importance_accumulator ≥ reflection_threshold (default 25). "
            "Fetches the 15 most-recent memories, sends them to the REFLECT prompt, receives 0-2 insights "
            "and 0-2 knowledge facts. Insights are stored as 'reflection' memories (importance=7). "
            "Novel knowledge lines are appended to the agent's KNOWLEDGE.md. "
            "The accumulator resets to 0 after reflection."
        )),
        ("maybe_plan()", (
            "Runs once at 08:00 each sim-day (daily_heartbeat resets current_plan to None). "
            "Fetches 5 recent reflections, sends them to the PLAN prompt, receives a JSON schedule "
            "{hour_08: '...', hour_09: '...', …, hour_17: '...'}. "
            "The plan is sanitised (invented NPCs removed, grounded), written to TODAY.md, "
            "stored in the plans table (importance=6)."
        )),
        ("propose_action()", (
            "Decays needs (hunger↑, energy↓, social_satiety↓). "
            "Retrieves 14 relevant memories via scored vector search. "
            "Builds a budgeted action prompt (≤12 000 chars) containing: bio, SOUL, "
            "character capsule, file context, needs, goals, sim_time, location, agents present, "
            "objects present, grounded action menu, plan chunk, recent memories. "
            "Sends to ACT prompt → receives {action, target, message, interaction, reasoning}. "
            "Applies anti-repetition and anti-stale-action guards. Returns a Proposal object."
        )),
        ("resolve() [ActionResolver]", (
            "Validates the proposal against WorldState. "
            "MoveProposal: target location must exist. "
            "SpeakProposal: target agent must be in same location, message must be non-empty. "
            "UseObjectProposal: object must be at this location, affordance must be in allowed list. "
            "On success: updates world state and emits an event. "
            "On failure: emits RejectedActionEvent with a reason string."
        )),
    ]
    for name, desc in steps:
        add_para(doc, name, bold=True)
        add_para(doc, desc)
        section_break(doc)

    # 4.2 Memory Retrieval
    add_heading(doc, "4.2  Memory Retrieval Scoring Algorithm", 2)
    add_para(doc, (
        "When an agent needs memories for an action prompt, the system scores all candidate memories "
        "from the database using a weighted formula:"
    ))
    add_code(doc, "score = 0.35 × recency + 0.25 × importance + 0.30 × relevance + 0.10 × metadata_bonus")
    section_break(doc)

    add_para(doc, "Component definitions:", bold=True)
    components = [
        ("recency", "Exponential decay: e^(−0.05 × ticks_ago).  Memories from the current tick score ~1.0; memories 20 ticks old score ~0.37."),
        ("importance", "Stored value (1–10) divided by 10.  A score of 7 → 0.70."),
        ("relevance", "Cosine similarity between the query embedding and the memory embedding.  Range 0–1."),
        ("metadata_bonus", "+0.15 if the memory's location tag matches the agent's current location; +0.15 if an agent in the memory matches a nearby agent.  Capped at 0.30 combined."),
    ]
    for name, desc in components:
        add_para(doc, f"• {name}: ", bold=False)
        p = doc.paragraphs[-1]
        p.runs[0].bold = True
        p.add_run(desc)

    section_break(doc)
    add_para(doc, "Candidate pool: the system fetches the 80 most-recent memories UNION the 80 highest-importance memories, deduplicates them, then scores all candidates.", italic=True)

    # 4.3 Action Resolution
    add_heading(doc, "4.3  Action Resolution Pipeline", 2)
    add_para(doc, "Proposals are sorted by priority before resolution:")
    priorities = [
        "Priority 1 (highest): SpeakProposal, UseObjectProposal",
        "Priority 2: MoveProposal",
        "Priority 3 (lowest): WaitProposal",
    ]
    for p in priorities:
        add_bullet(doc, p)

    section_break(doc)
    add_para(doc, "Resolution rules per proposal type:", bold=True)

    rules = {
        "MoveProposal": [
            "target_location must be a key in WorldState.locations",
            "On success: agent_positions[name] = target_location, emit MoveEvent",
            "On failure: emit RejectedActionEvent(reason='unknown location: …')",
        ],
        "SpeakProposal": [
            "target agent must be in the same location as the speaker",
            "message must be non-empty after stripping whitespace",
            "On success: emit SpeechEvent(speaker, listener, content, location)",
            "On failure: emit RejectedActionEvent",
        ],
        "UseObjectProposal": [
            "object must appear in WorldState.objects_at(current_location)",
            "interaction (affordance) must be in object's allowed_affordances list",
            "On success: update object state, emit ObjectStateChangeEvent",
            "On failure: emit RejectedActionEvent(reason='invalid affordance' or 'object not here')",
        ],
        "WaitProposal": [
            "Always succeeds — logs reason to transcript, emits no event.",
        ],
    }
    for kind, rule_list in rules.items():
        add_para(doc, kind, bold=True)
        for r in rule_list:
            add_bullet(doc, r, level=1)

    # 4.4 Importance scoring
    add_heading(doc, "4.4  Rule-Based Importance Scoring", 2)
    add_para(doc, "Importance is computed without an LLM call (src/agents/agent.py _score_importance):")
    add_code(doc, """\
base score  = 3
+ 2  if observation contains "said" or "moved from"
+ 1  if observation contains "rejected" or "changed from"
+ 1  if any other agent's name appears in the text
+ 1  if goal-relevant keywords appear ("goal", "plan", "promise",
         "argument", "help", "research", "writing")
clamped to range [1, 10]""")

    doc.add_page_break()

    # ── 5. LINE-BY-LINE CODE ANALYSIS ────────────────────────────────────────
    add_heading(doc, "5. Line-by-Line Code Analysis", 1)
    add_para(doc, (
        "This section walks through the most important source files. "
        "Line numbers match the current codebase. "
        "Comments starting with ► explain design decisions for junior developers."
    ))

    # 5.1 agent.py
    add_heading(doc, "5.1  src/agents/agent.py — Agent class", 2)

    agent_annotations = [
        ("1-12", "Imports",
         "Imports proposal types (data classes that carry agent decisions), "
         "the character capsule builder, AgentNeeds (tracks hunger/energy/social), "
         "CallKind (enum for LLM call routing), prompt builders, world events, and "
         "the action menu builder."),
        ("15-16", "Constants",
         "ACTION_PROMPT_CHAR_BUDGET = 12000 — the hard character limit for the action prompt. "
         "If the combined system+user prompt exceeds this, memories are dropped until it fits. "
         "MEMORY_TEXT_LIMIT = 240 — each memory snippet is truncated to 240 chars to save space."),
        ("19-57", "Agent.__init__",
         "Stores injected dependencies (memory, gateway, world, bus, transcript, repos). "
         "importance_accumulator tracks how much importance has built up since the last reflection. "
         "_recent_messages / _recent_reflections / _recent_locations / _recent_action_keys are "
         "short ring-buffers (max 8 items) used to detect and prevent repetition."),
        ("59-86", "perceive()",
         "① Finds current location and nearby agents/objects. "
         "② Fetches events from the EventBus for the previous tick. "
         "③ Calls _deterministic_observations() to build 3 text strings. "
         "④ Each observation is scored, stored in memory, and accumulates importance. "
         "► Perception is deterministic (no LLM) so it never hallucinates observations."),
        ("88-99", "_score_importance()",
         "Rule-based scoring described in §4.4. Runs in microseconds with no LLM call."),
        ("101-128", "maybe_reflect()",
         "Guards on importance_accumulator < threshold — skip if not enough has happened. "
         "Fetches 15 recent memories, compacts them (dedup + truncate), sends to REFLECT prompt. "
         "Stores insights as 'reflection' memories, appends novel knowledge to KNOWLEDGE.md. "
         "► Reflection is the system's 'thinking' step — it synthesises patterns from raw observations."),
        ("130-164", "maybe_plan()",
         "Runs once per sim-day (current_plan is reset to None at 08:00). "
         "Sends recent reflections + known world state to PLAN prompt. "
         "Sanitises the returned schedule (_sanitize_plan) to remove invented NPCs. "
         "Writes TODAY.md and stores plan in the plans table."),
        ("166-198", "propose_action()",
         "① Decays needs (called every tick). "
         "② Retrieves 14 memories most relevant to 'what should I do now?'. "
         "③ Builds a budgeted prompt (drops memories if over 12 000 chars). "
         "④ Calls ACT via LLMGateway → parses result → builds Proposal. "
         "⑤ _avoid_repetitive_speech() replaces a repeated speech with a move or object use. "
         "⑥ _avoid_stale_action() kicks the agent to the next plan location if it has lingered ≥2 ticks."),
        ("200-215", "_build_budgeted_action_prompt()",
         "Iteratively removes memories from the end of the list until the prompt fits the budget. "
         "If it still exceeds budget after removing all memories, truncates file_context to 2200 chars. "
         "► Ensures the LLM never receives a prompt that exceeds its context window."),
        ("217-224", "daily_heartbeat()",
         "Called every tick by SimulationEngine. Only acts at 08:00: resets TODAY.md, clears current_plan, "
         "then immediately calls maybe_plan(force=True) to generate the new day's schedule."),
        ("226-239", "end_of_day_heartbeat()",
         "Called every tick, only acts at 22:00. Fetches 20 recent memories, runs a REFLECT call, "
         "appends novel knowledge to KNOWLEDGE.md, and consolidates important facts into semantic_memory. "
         "► Semantic memory is a long-term store that persists across days with a confidence score."),
        ("241-256", "note_resolved_events()",
         "Called by SimulationEngine after each resolution. Checks if the agent spoke (→ socialize()), "
         "used food-related objects (→ satisfy_hunger(0.2)), or sat down (→ rest(0.15)). "
         "Persists updated needs to the database."),
        ("258-333", "_proposal_from_dict()",
         "Converts the raw LLM JSON dict into a typed Proposal. "
         "Handles the 'fallback' key (set by parsing.py when the LLM response is malformed). "
         "Gracefully degrades: if the intended target agent is not nearby, tries to move to them; "
         "if the object is not here, moves to its location. "
         "► Converts messy LLM output into grounded, validated actions."),
        ("335-353", "_deterministic_observations()",
         "Builds up to 3 observation strings from: fallback ('I am at X with Y'), "
         "recent events (speech→'Maria said to John: ...', move→'John moved from ... to ...', "
         "object change→'coffee_maker changed from idle to brewing'), and object states. "
         "► No LLM involved — all observations are facts from the world state."),
        ("431-537", "_normalize_object_target()",
         "A comprehensive alias map (e.g. 'coffee'→'coffee_maker', 'ledger'→'old_ledger') "
         "that translates LLM-generated target names to exact object IDs. "
         "Falls back to interaction-keyword matching if the alias map misses. "
         "► Prevents the LLM from breaking affordance checks by using informal object names."),
        ("539-598", "_normalize_interaction()",
         "Maps LLM-generated interaction verbs to valid affordance strings. "
         "e.g. 'typing' → 'write_article', 'archives' → 'search_records'. "
         "► Bridges the gap between natural language and the strict affordance vocabulary."),
        ("694-733", "_avoid_repetitive_speech() / _avoid_stale_action()",
         "_avoid_repetitive_speech(): uses Jaccard similarity on message terms to detect >52% overlap "
         "with recent messages; replaces with a move or object interaction. "
         "_avoid_stale_action(): if the agent has been at the same location for ≥2 ticks while the plan "
         "says to be elsewhere, it overrides the proposal with a MoveProposal. "
         "► These two guards are the primary mechanism preventing agents from getting stuck in loops."),
    ]

    for lines, title_text, explanation in agent_annotations:
        add_para(doc, f"Lines {lines} — {title_text}", bold=True)
        add_para(doc, explanation)
        section_break(doc)

    # 5.2 simulation.py
    add_heading(doc, "5.2  src/engine/simulation.py — SimulationEngine", 2)
    sim_annotations = [
        ("1-17", "Imports & constructor",
         "Holds references to all subsystems. agents is a list that grows as add_agent() is called. "
         "sim_tick starts at 0 and increments each tick."),
        ("23-55", "tick()",
         "The central orchestration method. "
         "① Increments sim_tick, formats sim_time (e.g. '09:30'). "
         "② daily_heartbeat() — each agent checks if it needs to re-plan. "
         "③ asyncio.gather(perceive) — all agents perceive in parallel (I/O-bound DB writes). "
         "④ Sequential reflect — reflections depend on accumulated importance so order matters. "
         "⑤ Sequential plan — each agent gets its own LLM call; parallelism would overload the local GPU. "
         "⑥ end_of_day_heartbeat() — consolidation runs sequentially at 22:00. "
         "⑦ asyncio.gather(propose_action) — proposals run in parallel (separate LLM calls). "
         "⑧ sorted(proposals, key=_resolution_priority) — Speak/Use go first so they can set world state "
         "before Move proposals potentially take agents away. "
         "⑨ For each resolved event, note_resolved_events() updates the acting agent's needs."),
        ("56-62", "_resolution_priority()",
         "Returns 1 (highest) for Speak/Use, 2 for Move, 3 (lowest) for Wait. "
         "► This ordering ensures conversational and object interactions happen before agents leave."),
        ("64-73", "run() / _format_time()",
         "run() is a simple loop calling tick() N times. "
         "_format_time() converts sim_tick to a clock string: "
         "tick 1 = 08:00, tick 2 = 08:30, …, tick 29 = 22:00, tick 48 = 07:30 (next day). "
         "Formula: total_min = (tick-1) × 30; hour = (8 + total_min // 60) % 24."),
    ]
    for lines, title_text, explanation in sim_annotations:
        add_para(doc, f"Lines {lines} — {title_text}", bold=True)
        add_para(doc, explanation)
        section_break(doc)

    # 5.3 memory.py (summary)
    add_heading(doc, "5.3  src/cognition/memory.py — MemoryService", 2)
    add_para(doc, (
        "MemoryService wraps the MemoryRepository and adds embedding-based retrieval. "
        "Key methods:"
    ))
    mem_methods = [
        ("add(agent_id, kind, content, importance, tick, time, metadata)",
         "Embeds the content, stores to DB. kind is one of: observation, reflection, plan, dialogue."),
        ("retrieve(agent_id, query, sim_tick, top_k, metadata_boosts)",
         "Fetches candidate memories (recent_n + important_n), embeds the query, "
         "scores each candidate with the formula in §4.2, returns top_k sorted by score."),
        ("add_semantic(agent_id, subject, fact, confidence, source_ids, sim_tick)",
         "Stores a durable fact (subject + fact text) in the semantic_memory table with a confidence score. "
         "Called at end-of-day consolidation."),
        ("get_semantic(agent_id, min_confidence)",
         "Returns durable facts above the confidence threshold, used in action prompts as 'semantic memory'."),
        ("mark_consolidated(source_ids)",
         "Sets the consolidated flag on memories that have been turned into semantic facts, "
         "preventing them from being used as retrieval candidates indefinitely."),
    ]
    for name, desc in mem_methods:
        add_para(doc, name, bold=True)
        add_para(doc, desc)
        section_break(doc)

    # 5.4 resolver.py (summary)
    add_heading(doc, "5.4  src/world/resolver.py — ActionResolver", 2)
    add_para(doc, (
        "ActionResolver.resolve(proposal) is called once per proposal per tick. "
        "It reads the current WorldState, validates the proposal, applies the change, "
        "emits events to the EventBus, and stores events in the database."
    ))
    add_para(doc, "Validation logic summary:", bold=True)
    add_code(doc, """\
MoveProposal:
  if target not in world.locations → RejectedActionEvent
  else → world.move_agent(name, target) → MoveEvent

SpeakProposal:
  if target not in world.agents_at(location) → RejectedActionEvent
  if message.strip() == "" → RejectedActionEvent
  else → SpeechEvent

UseObjectProposal:
  if object not in world.objects_at(location) → RejectedActionEvent
  if interaction not in world.object_affordances(object) → RejectedActionEvent
  else → world.set_object_state(object, new_state) → ObjectStateChangeEvent

WaitProposal:
  → no world-state change, no event emitted""")

    # 5.5 act.py (summary)
    add_heading(doc, "5.5  src/prompts/act.py — Action Prompt", 2)
    add_para(doc, (
        "act.build(bio, world_context, plan_chunk, memories, needs_string) returns "
        "(system_prompt, user_prompt). The system prompt contains:"
    ))
    prompt_parts = [
        "Agent bio (name, age, occupation, personality, goals)",
        "SOUL.md content (compact version, max 10 lines)",
        "Character capsule (personality summary + relationship notes for nearby agents)",
        "file_context (KNOWLEDGE.md + TODAY.md + semantic memory)",
        "Needs status string (e.g. 'hunger: moderate | energy: good | social: lonely')",
        "Goals list",
    ]
    for part in prompt_parts:
        add_bullet(doc, part)

    add_para(doc, "\nThe user prompt contains:", bold=True)
    user_parts = [
        "Current sim_time and location",
        "Agents present and objects here (exact IDs)",
        "Formatted action menu (valid actions with allowed targets)",
        "Current plan chunk (what the schedule says to do this hour)",
        "Recent activity summary (last 4 locations + action keys)",
        "Retrieved memories (truncated to MEMORY_TEXT_LIMIT each)",
    ]
    for part in user_parts:
        add_bullet(doc, part)

    add_para(doc, "\nExpected LLM response format:", bold=True)
    add_code(doc, '{"action": "move_to|speak_to|use_object|wait", "target": "...", "message": "...", "interaction": "...", "reasoning": "..."}')

    doc.add_page_break()

    # ── 6. FLOWCHARTS ─────────────────────────────────────────────────────────
    add_heading(doc, "6. Data Flow Diagrams (ASCII Flowcharts)", 1)

    add_heading(doc, "6.1  Full Tick Flow", 2)
    tick_flow = """\
START tick()
    │
    ├─► sim_tick += 1 ;  sim_time = _format_time()
    │
    ├─► [all agents] daily_heartbeat(tick, time)
    │       └─ if time == "08:00" and not already planned today:
    │               reset_today(), current_plan = None, maybe_plan(force=True)
    │
    ├─► [parallel] agent.perceive(tick, time)  ─── stores ≤3 observations in DB
    │       │
    │       └─ importance_accumulator += sum(importance of observations)
    │
    ├─► [sequential] agent.maybe_reflect(tick, time)
    │       └─ if accumulator ≥ threshold:
    │               LLM(REFLECT) → store insights, append KNOWLEDGE.md
    │               accumulator = 0
    │
    ├─► [sequential] agent.maybe_plan(tick, time)
    │       └─ if current_plan is None:
    │               LLM(PLAN) → sanitise → write TODAY.md → store in DB
    │
    ├─► [all agents] end_of_day_heartbeat(tick, time)
    │       └─ if time == "22:00" and not already done today:
    │               LLM(REFLECT) → append KNOWLEDGE.md → consolidate semantic
    │
    ├─► [parallel] proposal = agent.propose_action(tick, time)
    │       ├─ needs.decay_tick()
    │       ├─ retrieve 14 memories (scored)
    │       ├─ build budgeted prompt (trim until ≤ 12 000 chars)
    │       ├─ LLM(ACT) → parse JSON
    │       ├─ _proposal_from_dict() → typed Proposal
    │       ├─ _avoid_repetitive_speech()
    │       └─ _avoid_stale_action()
    │
    ├─► sort proposals: Speak/Use (1) > Move (2) > Wait (3)
    │
    ├─► [for each proposal] resolver.resolve(proposal)
    │       ├─ validate against WorldState
    │       ├─ apply world state change
    │       ├─ emit event to EventBus + DB + JSONL
    │       └─ on failure: RejectedActionEvent
    │
    └─► [for each event] acting_agent.note_resolved_events(events)
            └─ update needs (hunger/energy/social)

END tick()"""
    add_code(doc, tick_flow)

    add_heading(doc, "6.2  Memory Retrieval Flow", 2)
    mem_flow = """\
retrieve(agent_id, query, sim_tick, top_k=14)
    │
    ├─► embed(query)  →  query_vector  (float[N])
    │
    ├─► DB: SELECT * FROM memories WHERE agent_id=?
    │         ORDER BY sim_tick DESC  LIMIT 80          ← recent_n
    │       UNION
    │       SELECT * FROM memories WHERE agent_id=?
    │         ORDER BY importance DESC  LIMIT 80         ← important_n
    │
    ├─► deduplicate rows by id
    │
    ├─► for each candidate memory:
    │       recency   = e^(−0.05 × (sim_tick − memory.sim_tick))
    │       importance = memory.importance / 10
    │       relevance  = cosine_similarity(query_vector, memory.embedding)
    │       meta_bonus = 0
    │           + 0.15 if memory.location == current_location
    │           + 0.15 if memory.agents ∩ nearby_agents ≠ ∅
    │       score = 0.35×recency + 0.25×importance + 0.30×relevance + 0.10×meta_bonus
    │
    ├─► sort candidates by score DESC
    │
    └─► return top_k rows"""
    add_code(doc, mem_flow)

    add_heading(doc, "6.3  Proposal → Resolution Flow", 2)
    res_flow = """\
proposal = LLM output (one of: MoveProposal, SpeakProposal, UseObjectProposal, WaitProposal)
    │
    ▼
resolver.resolve(proposal)
    │
    ├─[MoveProposal]──────────────────────────────────────────────────────────┐
    │   target in world.locations?                                            │
    │       YES → world.move_agent(name, target)                             │
    │              emit MoveEvent(agent, from, to, tick, time)               │
    │       NO  → emit RejectedActionEvent(reason="unknown location")        │
    │                                                                         │
    ├─[SpeakProposal]────────────────────────────────────────────────────────┤
    │   target in world.agents_at(location)?  AND  message non-empty?        │
    │       YES → emit SpeechEvent(speaker, listener, content, location)     │
    │       NO  → emit RejectedActionEvent(reason="agent not nearby" or      │
    │                                                "empty message")         │
    │                                                                         │
    ├─[UseObjectProposal]────────────────────────────────────────────────────┤
    │   object in world.objects_at(location)?                                │
    │       NO  → RejectedActionEvent("object not here")                     │
    │   interaction in world.object_affordances(object)?                     │
    │       NO  → RejectedActionEvent("invalid affordance")                  │
    │       YES → world.set_object_state(object, new_state)                  │
    │              emit ObjectStateChangeEvent(object, old, new, location)   │
    │                                                                         │
    └─[WaitProposal]──────────────────────────────────────────────────────────┘
        → log reason to transcript, emit nothing

All events → EventBus.publish() → DB (event_repo.append()) → events.jsonl"""
    add_code(doc, res_flow)

    doc.add_page_break()

    # ── 7. DATABASE SCHEMA ────────────────────────────────────────────────────
    add_heading(doc, "7. Database Schema", 1)
    add_para(doc, "SQLite (default: simulation.sqlite3).  Postgres DDL also available in src/storage/postgres_schema.sql.")

    tables = {
        "agents": [
            ("id", "INTEGER PRIMARY KEY"),
            ("name", "TEXT UNIQUE — e.g. 'Maria'"),
            ("bio", "JSON TEXT — full bio dict from agents.yaml"),
            ("start_location", "TEXT — initial location"),
            ("state_json", "JSON TEXT — arbitrary runtime state"),
            ("needs_json", "JSON TEXT — {hunger, energy, social_satiety}"),
        ],
        "memories": [
            ("id", "INTEGER PRIMARY KEY"),
            ("agent_id", "INTEGER FK → agents.id"),
            ("kind", "TEXT — 'observation' | 'reflection' | 'plan' | 'dialogue'"),
            ("content", "TEXT — the memory text"),
            ("importance", "INTEGER 1–10"),
            ("embedding", "BLOB — serialised float vector"),
            ("sim_tick", "INTEGER — tick when stored"),
            ("sim_time", "TEXT — e.g. '09:30'"),
            ("metadata_json", "JSON — {location, agents, kind}"),
            ("consolidated", "INTEGER 0|1 — 1 after end-of-day consolidation"),
        ],
        "plans": [
            ("id", "INTEGER PRIMARY KEY"),
            ("agent_id", "INTEGER FK → agents.id"),
            ("schedule_json", "JSON TEXT — {hour_08: '...', …}"),
            ("sim_tick", "INTEGER"),
            ("sim_time", "TEXT"),
            ("sim_day", "INTEGER — tick // 48"),
        ],
        "relationships": [
            ("agent_id_1", "INTEGER"),
            ("agent_id_2", "INTEGER"),
            ("relationship_json", "JSON — {affinity, trust, familiarity, summary}"),
            ("PRIMARY KEY", "(agent_id_1, agent_id_2)"),
        ],
        "semantic_memory": [
            ("id", "INTEGER PRIMARY KEY"),
            ("agent_id", "INTEGER FK → agents.id"),
            ("subject", "TEXT — extracted subject (agent name, location, or self)"),
            ("fact", "TEXT — the durable fact"),
            ("confidence", "REAL 0.0–1.0"),
            ("source_ids", "JSON — list of memory IDs this was derived from"),
            ("sim_tick", "INTEGER"),
            ("sim_time", "TEXT"),
        ],
        "events": [
            ("id", "INTEGER PRIMARY KEY"),
            ("agent_name", "TEXT"),
            ("event_type", "TEXT — 'move' | 'speech' | 'object_state_change' | 'rejected'"),
            ("event_json", "JSON TEXT — full event payload"),
            ("sim_tick", "INTEGER"),
            ("sim_time", "TEXT"),
        ],
        "call_logs": [
            ("id", "INTEGER PRIMARY KEY"),
            ("kind", "TEXT — 'ACT' | 'PLAN' | 'REFLECT' | …"),
            ("agent_name", "TEXT"),
            ("prompt_tokens", "INTEGER"),
            ("response_tokens", "INTEGER"),
            ("duration_ms", "INTEGER"),
        ],
    }

    for table_name, cols in tables.items():
        add_para(doc, f"Table: {table_name}", bold=True)
        for col, desc in cols:
            add_bullet(doc, f"{col}  —  {desc}", level=1)
        section_break(doc)

    doc.add_page_break()

    # ── 8. CONFIGURATION REFERENCE ────────────────────────────────────────────
    add_heading(doc, "8. Configuration Reference", 1)

    add_para(doc, "config/simulation.yaml", bold=True)
    sim_cfg = [
        ("tick_duration_sim_seconds", "30", "Number of simulated seconds per tick (30 = 30 sim-minutes)"),
        ("reflection_threshold", "25", "Total importance needed before an agent reflects"),
        ("memory_recent_n", "80", "Number of most-recent memories in candidate pool"),
        ("memory_important_n", "80", "Number of highest-importance memories in candidate pool"),
    ]
    for key, default, desc in sim_cfg:
        add_bullet(doc, f"{key} (default: {default}) — {desc}")

    section_break(doc)
    add_para(doc, "config/agents.yaml — per-agent fields:", bold=True)
    agent_fields = [
        "name — display name",
        "age — used in bio string",
        "occupation — included in all prompts",
        "personality — freeform traits description",
        "goals — list of 3-5 goal strings",
        "preferences.places — preferred locations (for fresh-location fallback)",
        "traits — Big Five scores: openness, conscientiousness, extraversion, agreeableness, neuroticism",
        "needs — initial {hunger, energy, social_satiety} values (0.0–1.0)",
    ]
    for f in agent_fields:
        add_bullet(doc, f)

    section_break(doc)
    add_para(doc, "config/world.yaml — per-location fields:", bold=True)
    world_fields = [
        "description — 1-line description included in plan prompts",
        "objects — list of object IDs present",
    ]
    for f in world_fields:
        add_bullet(doc, f)

    add_para(doc, "\nPer-object fields:", bold=True)
    obj_fields = [
        "location — which location this object belongs to",
        "allowed_states — list of valid state strings",
        "default_state — initial state",
        "affordances — list of valid interaction verbs",
    ]
    for f in obj_fields:
        add_bullet(doc, f)

    doc.add_page_break()

    # ── 9. SOURCE CODE CHANGE HISTORY ─────────────────────────────────────────
    add_heading(doc, "9. Source Code Change History", 1)
    add_para(doc, (
        "This section summarises every significant commit so a junior developer can "
        "understand what changed, why, and which files were affected. "
        "Commits are listed newest-first."
    ))

    commits = [
        {
            "hash": "d4ee3de",
            "message": "agents: improve emergent action variety",
            "files": [
                "config/simulation.yaml — reflection_threshold tuned",
                "src/agents/agent.py — added _avoid_stale_action(), _avoid_repetitive_speech(), "
                "_object_use_would_be_noop(), _alternate_object_action(), _next_plan_location(), "
                "_fresh_location(), _has_lingered(); enriched _proposal_from_dict() fallback chains",
                "src/prompts/act.py — added recent_activity_summary to user prompt",
                "src/prompts/plan.py — minor prompt wording improvements",
                "src/prompts/reflect.py — tightened evidence-bound guardrails",
                "tests/test_agent_tick.py — 110 lines of new tests covering the new anti-loop guards",
            ],
            "why": (
                "Agents were getting stuck repeating the same object interaction or speech topic. "
                "This commit added runtime ring-buffers and guard methods to detect and break loops. "
                "The _avoid_stale_action() method checks if the agent has lingered ≥2 ticks and "
                "redirects it to the next planned location."
            ),
        },
        {
            "hash": "86db373",
            "message": "llm: verify AMD GPU offload",
            "files": [
                "README.md — added GPU setup section (Vulkan on AMD Radeon)",
                "src/llm/benchmark.py — added parse_gpu_offload_log(), verify_amd_gpu_offload(), "
                "start_llama_server(); reads llama.cpp startup logs to confirm layers offloaded to Vulkan0",
                "src/main.py — added --start-llama-server and --llama-ctx-size CLI flags",
                "tests/test_llm_benchmark.py — 45 lines of new tests",
            ],
            "why": (
                "The project runs on hardware with an AMD Radeon GPU. llama.cpp uses Vulkan for AMD. "
                "Without verification, a misconfigured server would silently run on CPU (10× slower). "
                "benchmark.py now parses the server startup log and raises an error if GPU offload "
                "is not confirmed."
            ),
        },
        {
            "hash": "5ceefae",
            "message": "world: expand cozy mystery town setting",
            "files": [
                "README.md — updated location list and setup guide",
                "config/agents.yaml — rewrote all three agent bios, goals, preferred places; "
                "aligned traits with the new cozy-mystery tone",
                "config/world.yaml — expanded from 1 location (cafe) to 10 locations; "
                "added 30+ objects each with affordances and allowed_states",
                "src/agents/agent.py — updated _fallback_plan() for all three agents; "
                "extended _normalize_object_target() alias map with all new objects; "
                "extended _normalize_interaction() keyword_map",
                "src/harness/agent_harness.py — updated DeterministicGateway actions for new world",
                "src/prompts/act.py — added known_world block to system prompt",
                "src/prompts/plan.py — added location_descriptions to plan prompt",
                "src/world/resolver.py — added object-location lookup for UseObjectProposal",
                "tests/test_agent_tick.py — 103 lines of new tests",
                "tests/test_world_context.py — 30 lines of new tests",
            ],
            "why": (
                "The original simulation only had a single cafe location. "
                "Expanding to 10 locations (library, archive, mill, riverside, market, etc.) "
                "enables agents to pursue distinct goals in different places, "
                "creating richer emergent interactions."
            ),
        },
        {
            "hash": "f49faca",
            "message": "reliability: harden simulation checks",
            "files": [
                ".env.example — documented all environment variables",
                "README.md — comprehensive setup guide with troubleshooting",
                "config/agents.yaml — added needs section per agent",
                "config/simulation.yaml — added memory pool sizes",
                "pyproject.toml — pinned key dependency versions",
                "src/agents/agent.py — large refactor: added needs integration, relationship absorption, "
                "semantic memory, budgeted prompt, daily/end-of-day heartbeats, file_context grounding",
                "src/agents/files.py — added reset_today(), improved load_context()",
                "src/agents/soul.py — new file: ensures SOUL.md exists on first run",
                "src/cognition/embeddings.py — abstracted EmbeddingModel interface",
                "src/cognition/memory.py — added scored retrieval, semantic memory methods",
                "src/cognition/needs.py — new file: AgentNeeds with hunger/energy/social",
                "src/engine/simulation.py — added end_of_day_heartbeat, note_resolved_events",
                "src/harness/__init__.py — exposed harness as importable package",
                "src/harness/agent_harness.py — new file: full deterministic harness (484 lines)",
                "src/llm/benchmark.py — new file: GPU verification (444 lines)",
                "src/llm/gateway.py — added call_log persistence, CallKind enum",
                "src/llm/parsing.py — added fallback parsing for malformed LLM responses",
                "src/main.py — full CLI rewrite with argparse",
                "src/prompts/act.py — added needs_string, file_context, action menu formatting",
                "src/prompts/plan.py — added known_world to plan prompt",
                "src/storage/db.py — added Postgres adapter",
                "src/storage/postgres_schema.sql — new file",
                "src/storage/repositories.py — full rewrite: 250 lines, all async with aiosqlite/asyncpg",
                "src/storage/schema.sql — added semantic_memory, events, call_logs tables",
                "tests/ — 9 new test files, 1000+ lines total",
            ],
            "why": (
                "This was the major reliability overhaul. Prior to this commit the simulation "
                "had no needs system, no semantic memory, no budgeted prompts, no GPU verification, "
                "and no deterministic test harness. Agents would frequently produce malformed LLM output "
                "with no graceful degradation. This commit addressed all of those gaps."
            ),
        },
        {
            "hash": "89c1255",
            "message": "merge: bring simulation project into main",
            "files": ["All src/ and config/ files — initial merge from feature branch"],
            "why": "First integration of the simulation into the main branch.",
        },
        {
            "hash": "3c43e46",
            "message": "test: cover simulation subsystems",
            "files": [
                "tests/__init__.py",
                "tests/test_agent_files.py — SOUL/KNOWLEDGE/TODAY.md read/write",
                "tests/test_agent_tick.py — perceive, reflect, plan, propose",
                "tests/test_llm_gateway.py — gateway call routing",
                "tests/test_memory.py — add + retrieve memories",
                "tests/test_personality.py — character capsule builder",
                "tests/test_repositories.py — DB CRUD operations",
                "tests/test_resolver.py — all four proposal types",
                "tests/test_simulation_clock.py — _format_time()",
                "tests/test_world_context.py — action menu builder",
            ],
            "why": "Established baseline test coverage before reliability work began.",
        },
    ]

    for commit in commits:
        add_para(doc, f"Commit {commit['hash']} — {commit['message']}", bold=True)
        add_para(doc, "Why this change was made:", bold=False, italic=True)
        add_para(doc, commit["why"])
        add_para(doc, "Files changed:", italic=True)
        for f in commit["files"]:
            add_bullet(doc, f, level=1)
        section_break(doc)

    doc.add_page_break()

    # ── 10. HOW TO RUN ────────────────────────────────────────────────────────
    add_heading(doc, "10. How to Run & Test", 1)

    add_heading(doc, "Prerequisites", 2)
    prereqs = [
        "Python 3.11+",
        "pip install -e . (installs all dependencies from pyproject.toml)",
        "llama.cpp server binary (optional: needed for real LLM calls)",
        "A GGUF model file, e.g. gemma-3-4b-it-Q4_K_M.gguf",
    ]
    for p in prereqs:
        add_bullet(doc, p)

    add_heading(doc, "Run modes", 2)
    run_modes = [
        ("Deterministic harness (no GPU needed — for CI/testing)",
         "python -m src.harness.agent_harness --ticks 8 --seed 42 --keep-artifacts",
         "Uses fake embeddings and hardcoded LLM responses. Output: .tmp/harness_report.json"),
        ("Full simulation (requires llama.cpp server on localhost:8080)",
         "python -m src.main --ticks 48 --start-llama-server --llama-ctx-size 16384",
         "Runs 48 ticks (1 sim-day). Output: simulation.sqlite3, events.jsonl, simulation.log, agents/"),
        ("Custom paths",
         "python -m src.main --ticks 5 --db .tmp/test.sqlite3 --events .tmp/events.jsonl --transcript .tmp/test.log --agents-root .tmp/agents --seed 42",
         "Useful for isolated test runs."),
        ("GPU benchmark",
         "python -m src.llm.benchmark",
         "Verifies AMD Vulkan offload. Diagnostic logs under .tmp/."),
        ("Unit tests",
         "python -m pytest -q",
         "Runs all tests in tests/. Expect ~60 tests in under 10 seconds."),
    ]
    for title_text, cmd, desc in run_modes:
        add_para(doc, title_text, bold=True)
        add_code(doc, cmd)
        add_para(doc, desc)
        section_break(doc)

    add_heading(doc, "Reading the outputs", 2)
    outputs = [
        ("simulation.sqlite3", "Open with DB Browser for SQLite or any SQLite tool. Query the memories table to see what each agent remembers."),
        ("events.jsonl", "One JSON object per line. Use jq or Python to filter. Each line has: agent_name, event_type, event_json, sim_tick, sim_time."),
        ("simulation.log", "Human-readable transcript. Search for 'REFLECT' to see insight moments, 'PLAN' for daily schedules, 'ACT' for decisions."),
        ("agents/*/KNOWLEDGE.md", "The agent's accumulated knowledge. Open in any text editor during a run to see it grow."),
        (".tmp/harness_report.json", "JSON report from the test harness. Contains pass/fail status and per-kind call metrics."),
    ]
    for name, desc in outputs:
        add_para(doc, name, bold=True)
        add_para(doc, desc)

    doc.add_page_break()

    # ── 11. GUARDRAILS ────────────────────────────────────────────────────────
    add_heading(doc, "11. Guardrails & Constraints", 1)
    add_para(doc, (
        "Agents are explicitly constrained to prevent hallucination and world-breaking behaviour. "
        "These constraints are enforced at two levels: prompt rules (soft) and code validation (hard)."
    ))

    add_para(doc, "Prompt-level rules (in act.py system prompt):", bold=True)
    prompt_rules = [
        "Only move to locations listed in 'known locations'",
        "Only speak to agents listed in 'agents present'",
        "Only use objects with affordances listed in the action menu",
        "Never invent people, NPCs, crimes, dates, or historical facts",
        "No 'Good morning' after 12:00; use 'Good afternoon' or 'Good evening'",
        "Use underscore_ids for targets (coffee_maker, not 'Coffee Maker')",
        "Don't repeat the same topic or visit the same object twice in a row",
    ]
    for r in prompt_rules:
        add_bullet(doc, r)

    section_break(doc)
    add_para(doc, "Code-level validation (hard rejections):", bold=True)
    code_rules = [
        "ActionResolver rejects moves to unknown locations",
        "ActionResolver rejects speech to agents not in the same location",
        "ActionResolver rejects object interactions with invalid affordances",
        "_normalize_object_target() maps informal names to exact IDs before validation",
        "_normalize_interaction() maps informal verbs to exact affordances before validation",
        "_sanitize_plan() removes invented NPC names and grounds historical claims",
        "_ground_historical_claims() replaces specific fabricated history with grounded terms",
        "_ground_agent_file_text() strips blocked terms from KNOWLEDGE.md before including in prompts",
    ]
    for r in code_rules:
        add_bullet(doc, r)

    doc.add_page_break()

    # ── 12. GLOSSARY ──────────────────────────────────────────────────────────
    add_heading(doc, "12. Glossary", 1)
    glossary = [
        ("Affordance", "A named action that a world object allows, e.g. 'brew_coffee' on the coffee_maker."),
        ("Agent", "A simulated character with memory, plans, and LLM-driven decisions."),
        ("Consolidation", "End-of-day process that turns episodic memories into durable semantic facts."),
        ("EventBus", "In-memory pub/sub system that delivers world events to agents at their location."),
        ("Grounded", "An action or claim is grounded if it refers only to things that exist in the world config."),
        ("Importance accumulator", "A running total of memory importance scores that triggers reflection when it exceeds the threshold."),
        ("KNOWLEDGE.md", "Append-only markdown file storing an agent's durable learned facts."),
        ("LLM", "Large Language Model — in this project, a local model served via llama.cpp."),
        ("Proposal", "A typed Python object (MoveProposal, SpeakProposal, etc.) returned by propose_action()."),
        ("Reflection", "An LLM-driven synthesis step that extracts insights and knowledge from recent memories."),
        ("Resolver", "ActionResolver — validates proposals against WorldState and applies changes."),
        ("Semantic memory", "Long-term (multi-day) storage of high-confidence facts with a subject + fact structure."),
        ("sim_tick", "An integer counter incremented each step. 1 tick = 30 simulated minutes."),
        ("SOUL.md", "A stable markdown file defining an agent's core personality and guardrails; never auto-edited."),
        ("TODAY.md", "The agent's daily schedule, reset and rewritten each sim-day at 08:00."),
        ("WorldState", "The in-memory record of all agent positions, object states, and location definitions."),
    ]
    for term, definition in glossary:
        p = doc.add_paragraph()
        p.add_run(term).bold = True
        p.add_run(f" — {definition}")

    # ── SAVE ──────────────────────────────────────────────────────────────────
    out_path = r"C:\Users\pssh\Documents\comunity Agent town\Generative_Agents_Project_Documentation_v2.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build_doc()
